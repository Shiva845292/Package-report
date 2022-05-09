from msilib import *
import os
import re
import docx
from datetime import date

today = date.today()
today=today.strftime("%d/%m/%Y")
print("Today's date:", today)
l=[]
for file in os.listdir():
    if (file.startswith("ZIC-")) or (file.startswith("ZCX-")) and (file.endswith(".mst")) :
        l.append(file)
    elif (file.startswith("ZIC-")) or (file.startswith("ZCX-")) and (file.endswith(".msi")) :
        l.append(file)
    elif (file.startswith("ZIC-")) or (file.startswith("ZCX-")) and (file.endswith(".exe")):
        l.append(file)
    else:
        pass
if len(l)==1 and l[0].endswith(".msi"):
    MSI=l[0]
    if (MSI.endswith(".msi")):
        def GetMsiProperty(path ,property):
            db = OpenDatabase(path, MSIDBOPEN_READONLY)
            view = db.OpenView ("SELECT Value FROM Property WHERE Property='" + property + "'")
            view.Execute(None)
            result = view.Fetch()
            #print dir(result)
            return result.GetString(1)
        def Summaryinfo(path,n):
            dbobject = OpenDatabase(path,MSIDBOPEN_READONLY)
            view = dbobject.GetSummaryInformation(200)
        # result = view.Fetch()
            return (view.GetProperty(n))

    path = os.path.join(os.getcwd(), MSI)
    Upgrde = GetMsiProperty(path, "UpgradeCode")
    ProductCode = GetMsiProperty(path, "ProductCode")
    Name = Summaryinfo(path, 4)
    UPN = Summaryinfo(path, 2)
    # Date = Summaryinfo(path, 6)

    def main():
        template_file_path = os.path.join(os.getcwd(), "Oldfile.docx")
        output_file_path = os.path.join(os.getcwd(), str(UPN.decode("utf-8")) + ".docx")

        variables = {
            "${EMPLOEE_NAME}": str(Name[0:17].decode("utf-8")),
            "${UPN}": str(UPN.decode("utf-8")),
            "${Ugrade code}": Upgrde,
            "${Product code}": ProductCode,
            "${Date}": today
        }

        template_document = docx.Document(template_file_path)

        for variable_key, variable_value in variables.items():
            for paragraph in template_document.paragraphs:
                replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        template_document.save(output_file_path)
        print('Docx is saved on the location...!')
        # print("UPN docx is Saved..!")


    def replace_text_in_paragraph(paragraph, key, value):
        if key in paragraph.text:
            inline = paragraph.runs
            for item in inline:
                if key in item.text:
                    item.text = item.text.replace(key, value)


    if __name__ == '__main__':
        main()


elif len(l)==1 :
     MST=l[0]
     if (MST.endswith(".mst")):
         for file in os.listdir():
             if file.endswith(".msi"):
                 MSI=file
                 # print("This MST",file)
                 def GetMsiProperty(path, property):
                     db = OpenDatabase(path, MSIDBOPEN_READONLY)
                     view = db.OpenView("SELECT Value FROM Property WHERE Property='" + property + "'")
                     view.Execute(None)
                     result = view.Fetch()
                     # print dir(result)
                     return result.GetString(1)

                 path = os.path.join(os.getcwd(), MSI)
                 Upgrde = GetMsiProperty(path, "UpgradeCode")
                 ProductCode = GetMsiProperty(path, "ProductCode")
                 # Name = Summaryinfo(path, 4)
                 n=len(MST)
                 UPN = MST[0:n - 4]
                 # print(type(UPN),UPN)
                 # Date = Summaryinfo(path, 6)

                 def main():
                     template_file_path = os.path.join(os.getcwd(), "Oldfile.docx")
                     output_file_path = os.path.join(os.getcwd(), UPN + ".docx")
                     print(MSI)
                     variables = {
                         "${EMPLOEE_NAME}": input("Enter Package name :"),
                         "{UPN}": str(UPN),
                         "${Ugrade code}": Upgrde,
                         "${Product code}": ProductCode,
                         "${Name}": str(MSI),
                         "${Date}": str(today),
                         "${MST}" : str(MST)
                     }

                     template_document = docx.Document(template_file_path)

                     for variable_key, variable_value in variables.items():
                         for paragraph in template_document.paragraphs:
                             replace_text_in_paragraph(paragraph, variable_key, variable_value)
                             # print(variable_value)

                         for table in template_document.tables:
                             for col in table.columns:
                                 for cell in col.cells:
                                     for paragraph in cell.paragraphs:
                                         replace_text_in_paragraph(paragraph, variable_key, variable_value)

                     template_document.save(output_file_path)
                     print('Docx is saved on the location...!')
                     # print("UPN docx is Saved..!")


                 def replace_text_in_paragraph(paragraph, key, value):
                     if key in paragraph.text:
                         inline = paragraph.runs
                         for item in inline:
                             if key in item.text:
                                 item.text = item.text.replace(key, value)
                                 # print(MS)


                 if __name__ == '__main__':
                     main()
elif len(l)==2:
    for i in l:
        print(i,end="=")



# print("Upgrade Code : ",Upgrde)
# print("Product Code : ",ProductCode)
# print("Name of engineer : ",Name)
# print("UPN Name of the package : ",UPN)
# print(Date[:10])

#
